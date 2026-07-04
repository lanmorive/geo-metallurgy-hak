"""Тесты ingest: noise filter, chunker, schemas."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.ingest.authors import extract_author_hint, fix_file_metadata_author
from app.ingest.chunker import _coalesce_blocks, blocks_to_chunks
from app.ingest.noise import is_noise
from app.ingest.parser import parse_file
from app.ingest.types import Block
from app.schemas.ontology import ParsedChunk


def test_is_noise_rejects_short_gibberish() -> None:
  assert is_noise("@@@ ### $$$")
  assert is_noise("short")
  assert not is_noise("Достаточно длинный осмысленный текст параграфа")


def test_is_noise_allows_short_table_cell() -> None:
  assert not is_noise("12,5", is_table_cell=True)


def test_extract_author_hint_doklad() -> None:
  assert extract_author_hint("Доклад_Румянцев А.Е.pdf") == "Румянцев А.Е"


def test_extract_author_hint_surname_initials() -> None:
  assert (
    extract_author_hint("Трофимов А.В. Опыт создания блочных геомеханических моделей в ГГИС.pdf")
    == "Трофимов А.В."
  )
  assert extract_author_hint("Тяпкина ПА_Пермь_Зимняя школа.pdf") == "Тяпкина ПА"


def test_extract_author_hint_english() -> None:
  assert extract_author_hint("Nicole_Roocke_report.pdf") == "Nicole Roocke"


def test_fix_pdf_metadata_author_garbage() -> None:
  garbage = "˜>4>E0=>20 !25B;048<8@>2=0"
  fixed = fix_file_metadata_author(garbage)
  assert fixed is None or any("\u0400" <= c <= "\u04FF" for c in fixed)


def test_fix_pdf_metadata_author_clean() -> None:
  assert fix_file_metadata_author("Julia Gershteyn") == "Julia Gershteyn"


def test_coalesce_slide_page() -> None:
  blocks = [
    Block(type="heading", text=f"Title {i}", page=1, section=f"s{i}")
    for i in range(10)
  ]
  coalesced = _coalesce_blocks(blocks)
  assert len(coalesced) == 1
  assert "Title 0" in coalesced[0].text
  assert "Title 9" in coalesced[0].text


def test_coalesce_short_blocks_merged() -> None:
  blocks = [
    Block(type="paragraph", text="A" * 20, page=1, section="intro"),
    Block(type="paragraph", text="B" * 30, page=1, section="intro"),
  ]
  coalesced = _coalesce_blocks(blocks)
  assert len(coalesced) == 1
  assert len(coalesced[0].text) == 51


def test_parsed_chunk_author_hint() -> None:
  blocks = [
    Block(
      type="paragraph",
      text="Достаточно длинный осмысленный текст параграфа для чанкинга.",
      section="intro",
    ),
  ]
  chunks = blocks_to_chunks(
    blocks,
    doc_id="doc_abc123",
    file_name="Доклад_Румянцев А.Е.pdf",
    source_key="raw/test.pdf",
    author_hint="Румянцев А.Е.",
  )
  assert len(chunks) == 1
  assert chunks[0].author_hint == "Румянцев А.Е."


def test_table_is_atomic_chunk() -> None:
  blocks = [
    Block(type="paragraph", text="Таблица 1 — химический состав руды", section="intro"),
    Block(
      type="table",
      text="| Cu | Fe |\n| --- | --- |\n| 0,35 | 12,4 |",
      section="intro",
    ),
    Block(type="paragraph", text="После таблицы идёт продолжение текста секции.", section="intro"),
  ]
  chunks = blocks_to_chunks(
    blocks,
    doc_id="doc_abc123",
    file_name="test.docx",
    source_key="raw/test.docx",
  )
  table_chunks = [c for c in chunks if c.kind == "table"]
  assert len(table_chunks) == 1
  assert "0,35" in table_chunks[0].text
  assert "Таблица 1" in table_chunks[0].text


def test_parsed_chunk_schema_roundtrip() -> None:
  chunk = ParsedChunk(
    doc_id="doc_abc",
    chunk_id="doc_abc_00001",
    text="sample",
    kind="text",
    section="frontmatter",
    page=None,
    lang="ru",
    file_name="f.docx",
    source_key="raw/f.docx",
    author_hint="Иванов И.И.",
  )
  restored = ParsedChunk.model_validate_json(chunk.model_dump_json())
  assert restored == chunk


def test_ocr_skipped_when_tesseract_missing(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
  import fitz

  from app.ingest import pdf_parser

  pdf_path = tmp_path / "blank.pdf"
  doc = fitz.open()
  doc.new_page()
  doc.save(str(pdf_path))
  doc.close()

  monkeypatch.setattr(pdf_parser, "_tesseract_available", lambda: False)
  result = pdf_parser.parse_pdf(pdf_path)
  assert result.doc_meta.ocr_pages == 0
  assert result.blocks == []


def test_parse_docx_with_table(tmp_path: Path) -> None:
  path = tmp_path / "sample.docx"
  doc = Document()
  doc.add_heading("Введение", level=1)
  doc.add_paragraph("Аннотация документа с достаточным количеством слов для фильтра.")
  table = doc.add_table(rows=2, cols=2)
  table.rows[0].cells[0].text = "Cu, %"
  table.rows[0].cells[1].text = "Ni, %"
  table.rows[1].cells[0].text = "0,35"
  table.rows[1].cells[1].text = "1,2"
  doc.save(str(path))

  result, noise, refs = parse_file(path)
  assert noise >= 0
  assert not refs
  assert any(b.type == "table" for b in result.blocks)
  table_block = next(b for b in result.blocks if b.type == "table")
  assert "0,35" in table_block.text

  chunks = blocks_to_chunks(
    result.blocks,
    doc_id="doc_test",
    file_name="sample.docx",
    source_key="raw/sample.docx",
  )
  assert chunks
  for chunk in chunks:
    ParsedChunk.model_validate(chunk)
